"""PDF and document text extraction service."""
import os
from pathlib import Path
from typing import Optional, Tuple
from loguru import logger

from ..core.config import settings
from ..core.exceptions import DocumentProcessingError


class PDFReader:
    """Extract text from PDF, DOCX, and TXT files."""

    def extract_text(self, file_path: str) -> Tuple[str, bool]:
        """
        Extract text from a file.
        Returns (extracted_text, used_ocr).
        """
        path = Path(file_path)
        if not path.exists():
            raise DocumentProcessingError(f"File not found: {file_path}")

        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._extract_from_pdf(file_path)
        elif ext == ".docx":
            return self._extract_from_docx(file_path)
        elif ext == ".txt":
            return self._extract_from_txt(file_path)
        elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            return self._extract_from_image(file_path)
        else:
            raise DocumentProcessingError(f"Unsupported file type: {ext}")

    def _extract_from_pdf(self, file_path: str) -> Tuple[str, bool]:
        """Extract text from PDF. Falls back to OCR if needed."""
        text = ""
        used_ocr = False

        # Try pdfplumber first (better for text-based PDFs)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")

        # If minimal text extracted, try PyPDF2
        if len(text.strip()) < 50:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {e}")

        # If still minimal text, use OCR on PDF pages
        if len(text.strip()) < 50:
            logger.info(f"Minimal text extracted, attempting OCR on PDF: {file_path}")
            text = self._ocr_pdf(file_path)
            used_ocr = True

        return text.strip(), used_ocr

    def _ocr_pdf(self, file_path: str) -> str:
        """Perform OCR on a PDF file."""
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(file_path, dpi=300)
            text = ""
            ocr_engine = OCREngine()
            for i, image in enumerate(images):
                image_path = f"/tmp/pdf_page_{i}.png"
                image.save(image_path, "PNG")
                page_text = ocr_engine.extract_text(image_path)
                text += page_text + "\n"
            return text
        except ImportError:
            logger.warning("pdf2image not installed, trying direct OCR")
            return self._direct_pdf_ocr(file_path)
        except Exception as e:
            logger.error(f"PDF OCR failed: {e}")
            raise DocumentProcessingError(f"PDF OCR failed: {str(e)}")

    def _direct_pdf_ocr(self, file_path: str) -> str:
        """Direct OCR on PDF using pytesseract."""
        try:
            import pytesseract
            from pdf2image import convert_from_path
            images = convert_from_path(file_path)
            text = ""
            for image in images:
                page_text = pytesseract.image_to_string(image, lang=settings.OCR_LANGUAGE)
                text += page_text + "\n"
            return text
        except Exception as e:
            raise DocumentProcessingError(f"Direct PDF OCR failed: {str(e)}")

    def _extract_from_docx(self, file_path: str) -> Tuple[str, bool]:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    text += "\n" + row_text

            return text.strip(), False
        except Exception as e:
            raise DocumentProcessingError(f"DOCX extraction failed: {str(e)}")

    def _extract_from_txt(self, file_path: str) -> Tuple[str, bool]:
        """Read text from TXT file."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            return text.strip(), False
        except Exception as e:
            raise DocumentProcessingError(f"TXT read failed: {str(e)}")

    def _extract_from_image(self, file_path: str) -> Tuple[str, bool]:
        """Extract text from image using OCR."""
        ocr_engine = OCREngine()
        text = ocr_engine.extract_text(file_path)
        return text.strip(), True


class OCREngine:
    """OCR engine for extracting text from images."""

    def __init__(self):
        self._easyocr_reader = None

    def extract_text(self, image_path: str) -> str:
        """
        Extract text from an image.
        Tries pytesseract first, falls back to EasyOCR.
        """
        # Try pytesseract
        try:
            import pytesseract
            from PIL import Image

            img = Image.open(image_path)
            text = pytesseract.image_to_string(img, lang=settings.OCR_LANGUAGE)
            if len(text.strip()) > 20:
                return text
        except ImportError:
            logger.warning("pytesseract not available, trying EasyOCR")
        except Exception as e:
            logger.warning(f"Tesseract OCR failed: {e}")

        # Fall back to EasyOCR
        try:
            return self._easyocr_extract(image_path)
        except Exception as e:
            raise DocumentProcessingError(f"All OCR methods failed: {str(e)}")

    def _easyocr_extract(self, image_path: str) -> str:
        """Extract text using EasyOCR."""
        if self._easyocr_reader is None:
            import easyocr
            self._easyocr_reader = easyocr.Reader(["en"], gpu=False)

        results = self._easyocr_reader.readtext(image_path)
        text = " ".join([detection[1] for detection in results])
        return text


# Singleton instances
pdf_reader = PDFReader()
ocr_engine = OCREngine()
